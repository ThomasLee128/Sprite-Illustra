"""
文件名：docx_parser.py
功能描述：Word 文档（.docx）解析器。使用 python-docx 库提取文档结构，
         识别标题、段落、列表等元素，生成统一的 Document 结构。
作者：Claude Code
创建时间：2026-04-04
后续开发：Trea
TODO：
- [x] TODO-1: 实现 parse(self, file_path: str) -> Document

依赖：python-docx (from docx import Document as DocxDocument), schemas.document, core.exceptions, uuid
"""

import uuid
from pathlib import Path

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from parsers.base import BaseParser
from parsers import register_parser
from schemas.document import Document, DocumentSection, DocumentMetadata, SectionType
from core.exceptions import DocumentParseError


@register_parser
class DocxParser(BaseParser):
    """Word 文档解析器"""

    supported_extensions = [".docx"]

    async def parse(self, file_path: str) -> Document:
        """解析 .docx 文件"""
        if not HAS_DOCX:
            raise DocumentParseError("需要安装 python-docx 库才能解析 Word 文档")
        
        try:
            doc = DocxDocument(file_path)
            sections: list[DocumentSection] = []
            position = 0
            first_heading = ""
            total_words = 0
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                total_words += len(text)
                
                # 判断段落类型
                section_type = SectionType.PARAGRAPH
                level = 0
                
                style_name = para.style.name.lower() if para.style else ""
                
                if style_name.startswith("heading"):
                    section_type = SectionType.HEADING
                    # 提取标题级别
                    import re
                    match = re.search(r"heading\s*(\d+)", style_name, re.IGNORECASE)
                    if match:
                        level = int(match.group(1))
                    else:
                        level = 1
                elif style_name.startswith("list"):
                    section_type = SectionType.LIST
                elif para.style and para.style.name in ["List Bullet", "List Number"]:
                    section_type = SectionType.LIST
                
                section = DocumentSection(
                    id=str(uuid.uuid4()),
                    type=section_type,
                    level=level,
                    content=text,
                    position=position,
                )
                sections.append(section)
                
                if not first_heading and section_type == SectionType.HEADING:
                    first_heading = text
                
                position += 1
            
            # 提取元数据
            filename = Path(file_path).name
            title = first_heading or filename
            if doc.core_properties.title:
                title = doc.core_properties.title
            
            metadata = DocumentMetadata(
                title=title,
                author=doc.core_properties.author or "",
                word_count=total_words,
            )
            
            return Document(
                filename=filename,
                source_format="docx",
                sections=sections,
                metadata=metadata,
            )
            
        except Exception as e:
            raise DocumentParseError(f"解析 Word 文档失败: {str(e)}")
