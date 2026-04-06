"""
文件名：word_exporter.py
功能描述：Word (.docx) 导出器。将 IllustratedDocument 导出为 Word 文档，
         在对应段落后插入 AI 生成的配图。
作者：Claude Code
创建时间：2026-04-04
后续开发：Trea
TODO：
- [x] TODO-1: 实现 export(self, illustrated_doc, output_path) -> str

依赖：python-docx (from docx import Document as DocxDocument, from docx.shared import Inches)
      schemas.document, core.exceptions
"""

from pathlib import Path

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from exporters.base import BaseExporter
from exporters import register_exporter
from schemas.document import IllustratedDocument, SectionType, IllustrationStatus
from core.exceptions import ExportError


@register_exporter
class WordExporter(BaseExporter):
    """Word 文档导出器"""

    format_name = "docx"
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_extension = ".docx"

    async def export(self, illustrated_doc: IllustratedDocument, output_path: str) -> str:
        """导出为 Word 文档"""
        import logging
        logger = logging.getLogger(__name__)
        
        if not HAS_DOCX:
            raise ExportError("需要安装 python-docx 库才能导出 Word 文档")
        
        logger.info(f"[Word 导出器] 开始导出，共有 {len(illustrated_doc.illustrations)} 个插图")
        
        doc = DocxDocument()
        
        # 创建 section.id -> section 的映射
        section_map = {s.id: s for s in illustrated_doc.document.sections}
        
        # 创建 after_section_id -> illustrations 的映射
        illustrations_map: dict[str, list] = {}
        for illu in illustrated_doc.illustrations:
            logger.info(f"[Word 导出器] 检查插图: {illu.id}, status={illu.status}, image_path={illu.image_path}")
            if illu.status != IllustrationStatus.DONE or not illu.image_path:
                continue
            if illu.after_section_id not in illustrations_map:
                illustrations_map[illu.after_section_id] = []
            illustrations_map[illu.after_section_id].append(illu)
        
        logger.info(f"[Word 导出器] 找到 {len(illustrations_map)} 个section有插图需要插入")
        
        # 按顺序输出 sections 和插图
        for section in illustrated_doc.document.sections:
            # 输出 section
            if section.type == SectionType.HEADING:
                level = min(section.level, 6)
                doc.add_heading(section.content, level=level)
            elif section.type == SectionType.PARAGRAPH:
                doc.add_paragraph(section.content)
            elif section.type == SectionType.LIST:
                p = doc.add_paragraph(section.content)
                p.style = 'List Bullet'
            elif section.type == SectionType.CODE:
                p = doc.add_paragraph(section.content)
                # 设置等宽字体
                for run in p.runs:
                    run.font.name = 'Courier New'
            elif section.type == SectionType.BLOCKQUOTE:
                p = doc.add_paragraph(section.content)
                # 设置缩进
                p.paragraph_format.left_indent = Inches(0.5)
            
            # 输出该 section 后的插图
            if section.id in illustrations_map:
                for illu in illustrations_map[section.id]:
                    src_path = Path(illu.image_path)
                    logger.info(f"[Word 导出器] 尝试插入图片: {src_path}, 存在={src_path.exists()}")
                    if src_path.exists():
                        try:
                            # 添加图片
                            doc.add_picture(str(src_path), width=Inches(5))
                            # 添加图注
                            if illu.description_cn:
                                caption = doc.add_paragraph(illu.description_cn)
                                caption.alignment = 1  # 居中
                            logger.info(f"[Word 导出器] 图片插入成功: {src_path}")
                        except Exception as e:
                            logger.error(f"[Word 导出器] 图片插入失败: {e}")
                            # 跳过图片添加失败的情况
                            pass
                    else:
                        logger.warning(f"[Word 导出器] 图片文件不存在: {src_path}")
        
        # 保存文档
        doc.save(output_path)
        
        return output_path
